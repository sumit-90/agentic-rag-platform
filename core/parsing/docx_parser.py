import time
from pathlib import Path

import structlog
from docx import Document

from core.parsing.base_parser import BaseParser, ParsedDocument

logger = structlog.get_logger(__name__)


class DOCXParser(BaseParser):
    def parse(self, file_path: Path) -> ParsedDocument:
        log = logger.bind(file_path=str(file_path))

        if not file_path.exists():
            log.error("docx_file_not_found")
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        file_size_bytes = file_path.stat().st_size
        log = log.bind(file_size_bytes=file_size_bytes)
        log.info("docx_parse_start")

        start = time.perf_counter()

        try:
            doc = Document(str(file_path))
        except Exception as exc:
            log.warning("docx_open_failed", error=str(exc))
            raise ValueError(f"Cannot open DOCX '{file_path}': {exc}") from exc

        try:
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n".join(paragraphs)
            paragraph_count = len(paragraphs)
        except Exception as exc:
            log.warning("docx_text_extraction_failed", error=str(exc))
            raise ValueError(f"Failed to extract text from DOCX '{file_path}': {exc}") from exc

        duration = time.perf_counter() - start
        log.info("docx_parse_complete", paragraph_count=paragraph_count, duration_s=round(duration, 4))

        return ParsedDocument(
            content=content,
            metadata={
                "file_type": "docx",
                "file_size_bytes": file_size_bytes,
                "paragraph_count": paragraph_count,
            },
            source=str(file_path),
        )
